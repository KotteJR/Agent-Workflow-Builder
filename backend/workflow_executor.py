"""
Graph-Based Workflow Executor

Executes user-defined workflows respecting graph topology.
Only executes nodes that are reachable from input nodes,
following strict dependency order using topological sort.
"""

import json
import time
import uuid
from collections import deque
from typing import Any, AsyncGenerator, Dict, List, Optional, Set, Tuple

from config import config
from models import get_llm_client, LLMClientProtocol
from agents.base import AgentResult
from agents.supervisor import SupervisorAgent
from agents.orchestrator import OrchestratorAgent
from agents.semantic_search import SemanticSearchAgent
from agents.sampler import SamplerAgent
from agents.synthesis import SynthesisAgent
from agents.summarization import SummarizationAgent
from agents.formatting import FormattingAgent
from agents.transformer import TransformerAgent
from agents.image_generator import ImageGeneratorAgent
from agents.translator import TranslatorAgent
import retrieval
from demo_handler import is_demo_workflow
from workflow_logger import debugger, workflow_logger


# Agent registry - maps node types to agent classes
AGENT_REGISTRY = {
    "supervisor": SupervisorAgent,
    "orchestrator": OrchestratorAgent,
    "semantic_search": SemanticSearchAgent,
    "sampler": SamplerAgent,
    "synthesis": SynthesisAgent,
    "summarization": SummarizationAgent,
    "formatting": FormattingAgent,
    "transformer": TransformerAgent,
    "image_generator": ImageGeneratorAgent,
    "translator": TranslatorAgent,
}

# Input node types (don't require agents)
INPUT_NODE_TYPES = {"prompt", "upload"}

# Output node types (don't require agents)
OUTPUT_NODE_TYPES = {"response", "spreadsheet"}


def topological_sort(nodes: List[str], edges: List[Dict[str, str]]) -> List[str]:
    """
    Topologically sort nodes based on edges.
    Returns nodes in execution order (dependencies first).
    """
    # Build in-degree count
    in_degree: Dict[str, int] = {node: 0 for node in nodes}
    adjacency: Dict[str, List[str]] = {node: [] for node in nodes}
    
    for edge in edges:
        source = edge["source"]
        target = edge["target"]
        if source in nodes and target in nodes:
            adjacency[source].append(target)
            in_degree[target] = in_degree.get(target, 0) + 1
    
    # Kahn's algorithm
    queue = deque([node for node in nodes if in_degree[node] == 0])
    result = []
    
    while queue:
        node = queue.popleft()
        result.append(node)
        
        for neighbor in adjacency.get(node, []):
            in_degree[neighbor] -= 1
            if in_degree[neighbor] == 0:
                queue.append(neighbor)
    
    # Add any remaining nodes (cycles or disconnected)
    remaining = [node for node in nodes if node not in result]
    result.extend(remaining)
    
    return result


def find_reachable_nodes(
    start_nodes: Set[str],
    edges: List[Dict[str, str]],
    all_nodes: Set[str],
) -> Set[str]:
    """
    Find all nodes reachable from start nodes via BFS.
    This ensures only connected nodes are executed.
    """
    # Build adjacency list (forward direction)
    adjacency: Dict[str, Set[str]] = {node: set() for node in all_nodes}
    for edge in edges:
        if edge["source"] in all_nodes and edge["target"] in all_nodes:
            adjacency[edge["source"]].add(edge["target"])
    
    # BFS from start nodes
    reachable = set(start_nodes)
    queue = deque(start_nodes)
    
    while queue:
        node = queue.popleft()
        for neighbor in adjacency.get(node, set()):
            if neighbor not in reachable:
                reachable.add(neighbor)
                queue.append(neighbor)
    
    return reachable


def get_node_dependencies(node_id: str, edges: List[Dict[str, str]]) -> List[str]:
    """Get all nodes that must execute before this node."""
    return [edge["source"] for edge in edges if edge["target"] == node_id]


def _sse_event(event_type: str, data: Dict[str, Any]) -> str:
    """Format a Server-Sent Event."""
    return f"event: {event_type}\ndata: {json.dumps(data)}\n\n"


async def execute_workflow(
    user_message: str,
    workflow_nodes: List[Dict[str, Any]],
    workflow_edges: List[Dict[str, str]],
) -> AsyncGenerator[str, None]:
    """
    Execute a custom workflow defined by nodes and edges.
    
    Respects graph topology:
    - Only executes nodes reachable from input nodes
    - Follows topological order (dependencies first)
    - Passes context between connected nodes
    
    Args:
        user_message: User's query
        workflow_nodes: List of node configurations
        workflow_edges: List of edges defining connections
        
    Yields:
        SSE events for workflow execution
    """
    # Initialize debugging session
    execution_id = str(uuid.uuid4())[:8]
    debugger.start_execution(execution_id)
    
    workflow_start = time.time()
    steps: List[Dict[str, Any]] = []
    
    # Get LLM client and models
    llm_client = get_llm_client()
    small_model = config.get_model_config()["small"]
    large_model = config.get_model_config()["large"]
    
    workflow_logger.debug(f"Using models - small: {small_model}, large: {large_model}")
    
    # Extract node IDs and types
    node_map = {node["id"]: node for node in workflow_nodes}
    all_node_ids = set(node_map.keys())
    
    # Log node types for clarity
    workflow_logger.debug("Node registry:")
    for nid, node in node_map.items():
        ntype = node.get("data", {}).get("nodeType", "unknown")
        workflow_logger.debug(f"  {nid} -> {ntype}")
    
    # Find input nodes (nodes with no incoming edges or input types)
    nodes_with_incoming = {edge["target"] for edge in workflow_edges}
    input_nodes = set()
    
    for node_id, node in node_map.items():
        node_type = node.get("data", {}).get("nodeType", node_id.split("-")[0])
        if node_type in INPUT_NODE_TYPES or node_id not in nodes_with_incoming:
            input_nodes.add(node_id)
    
    # Find all nodes reachable from inputs
    reachable_nodes = find_reachable_nodes(input_nodes, workflow_edges, all_node_ids)
    
    # Filter edges to only include reachable nodes
    valid_edges = [
        edge for edge in workflow_edges
        if edge["source"] in reachable_nodes and edge["target"] in reachable_nodes
    ]
    
    # Topologically sort reachable nodes
    execution_order = topological_sort(list(reachable_nodes), valid_edges)
    
    # Log workflow setup
    debugger.log_workflow_setup(input_nodes, reachable_nodes, execution_order, valid_edges)
    
    # Extract spreadsheet node settings if present (for transformer to use)
    spreadsheet_settings = {}
    has_spreadsheet_output = False
    for node_id in reachable_nodes:
        node = node_map.get(node_id)
        if node:
            node_data = node.get("data", {})
            if node_data.get("nodeType") == "spreadsheet":
                spreadsheet_settings = node_data.get("settings", {})
                has_spreadsheet_output = True
                print(f"[WORKFLOW] Found spreadsheet settings: {spreadsheet_settings}")
                break
    
    # === DEMO MODE CHECK ===
    # Check if this is a demo workflow with premade output
    uploaded_files = []
    for node_id in input_nodes:
        node = node_map.get(node_id)
        if node:
            node_data = node.get("data", {})
            if node_data.get("nodeType") == "upload":
                uploaded_files = node_data.get("uploadedFiles", [])
                break
    
    demo_output = is_demo_workflow(uploaded_files, has_spreadsheet_output)
    if demo_output:
        print(f"[DEMO] Using premade demo output for this workflow")
        workflow_latency = round((time.time() - workflow_start) * 1000, 2)
        
        # Yield demo execution events
        yield _sse_event("agent_start", {"agent": "demo", "status": "working"})
        yield _sse_event("agent_complete", {
            "agent": "demo",
            "step": {
                "agent": "demo",
                "model": "demo-mode",
                "action": "demo_extraction",
                "content": "Using optimized demo extraction for this document",
            }
        })
        
        # Return the premade output
        yield _sse_event("done", {
            "answer": demo_output,
            "tool_outputs": {"images": [], "calculations": [], "web_results": [], "docs": []},
            "trace": {"steps": [{"agent": "demo", "model": "demo-mode", "action": "demo_extraction", "content": demo_output}]},
            "latency_ms": workflow_latency,
            "output_format": "spreadsheet",
        })
        return
    # === END DEMO MODE CHECK ===
    
    # Execution context - shared state between nodes
    context: Dict[str, Any] = {
        "user_message": user_message,
        "context_snippets": [],
        "candidates": [],
        "semantic_results": [],
        "docs": [],
        "tool_outputs": {
            "images": [],
            "calculations": [],
            "web_results": [],
        },
        "orchestrator_result": {"tools_to_execute": []},
        "final_answer": "",
        "spreadsheet_settings": spreadsheet_settings,  # Pass to transformer
    }
    
    # Track executed and excluded nodes
    executed_nodes: Set[str] = set()
    excluded_nodes: Set[str] = set()
    
    try:
        for node_id in execution_order:
            node = node_map.get(node_id)
            if not node:
                workflow_logger.warning(f"Node {node_id} not found in node_map - skipping")
                continue
            
            node_data = node.get("data", {})
            node_type = node_data.get("nodeType", node_id.split("-")[0])
            node_settings = node_data.get("settings", {})
            
            # Get dependencies for this node
            dependencies = get_node_dependencies(node_id, valid_edges)
            
            # Log node evaluation start
            debugger.log_node_start(node_id, node_type, dependencies)
            debugger.log_dependency_check(node_id, dependencies, executed_nodes, excluded_nodes)
            
            # Handle input nodes
            if node_type in INPUT_NODE_TYPES:
                executed_nodes.add(node_id)
                
                # Extract content from input nodes
                if node_type == "upload":
                    # Get uploaded files from node data
                    uploaded_files = node_data.get("uploadedFiles", [])
                    print(f"[UPLOAD] ========================================")
                    print(f"[UPLOAD] Processing upload node: {node_id}")
                    print(f"[UPLOAD] Node data keys: {list(node_data.keys())}")
                    print(f"[UPLOAD] Found {len(uploaded_files)} uploaded files")
                    if uploaded_files:
                        # Extract file content for context
                        file_contents = []
                        for file_info in uploaded_files:
                            file_name = file_info.get("name", "unknown")
                            file_content = file_info.get("content", "")
                            
                            print(f"[UPLOAD] File: {file_name}, content length: {len(file_content) if file_content else 0}")
                            print(f"[UPLOAD] Content starts with: {file_content[:50] if file_content else 'NONE'}...")
                            
                            if file_content:
                                # Parse PDF files
                                if file_content.startswith("__PDF_BASE64__"):
                                    try:
                                        import base64
                                        from io import BytesIO
                                        from pypdf import PdfReader
                                        
                                        pdf_base64 = file_content[14:]  # Remove prefix
                                        pdf_bytes = base64.b64decode(pdf_base64)
                                        pdf_reader = PdfReader(BytesIO(pdf_bytes))
                                        
                                        # Step 1: Try to extract text directly (works for text-based PDFs)
                                        text_parts = []
                                        for i, page in enumerate(pdf_reader.pages):
                                            page_text = page.extract_text()
                                            if page_text:
                                                text_parts.append(f"[Page {i+1}]\n{page_text}")
                                        
                                        extracted_text = "\n\n".join(text_parts)
                                        
                                        # Step 2: If very little text extracted, it's likely a scanned PDF - use OCR
                                        if len(extracted_text.strip()) < 100:  # Threshold for scanned PDF detection
                                            print(f"[UPLOAD] PDF appears to be scanned (only {len(extracted_text)} chars extracted), attempting OCR...")
                                            ocr_success = False
                                            try:
                                                from pdf2image import convert_from_bytes
                                                import pytesseract
                                                
                                                print(f"[UPLOAD] OCR libraries found. Converting PDF to images...")
                                                # Convert PDF pages to images
                                                images = convert_from_bytes(pdf_bytes, dpi=300)  # Higher DPI for better OCR
                                                print(f"[UPLOAD] Converted to {len(images)} page images. Running OCR...")
                                                
                                                # Extract text from each page using OCR
                                                # Use multiple languages: English + Arabic for best coverage
                                                ocr_langs = 'eng+ara'  # Supports mixed English/Arabic documents
                                                ocr_text_parts = []
                                                for i, image in enumerate(images):
                                                    print(f"[UPLOAD] Running OCR on page {i+1}/{len(images)} (langs: {ocr_langs})...")
                                                    page_ocr_text = pytesseract.image_to_string(image, lang=ocr_langs)
                                                    if page_ocr_text.strip():
                                                        ocr_text_parts.append(f"[Page {i+1} - OCR]\n{page_ocr_text}")
                                                        print(f"[UPLOAD] Page {i+1}: Extracted {len(page_ocr_text)} chars via OCR")
                                                
                                                ocr_text = "\n\n".join(ocr_text_parts)
                                                
                                                if ocr_text.strip():
                                                    extracted_text = ocr_text
                                                    ocr_success = True
                                                    print(f"[UPLOAD] ✅ OCR SUCCESS: Extracted {len(extracted_text)} chars from scanned PDF")
                                                else:
                                                    print(f"[UPLOAD] ⚠️ OCR completed but extracted no text")
                                                    
                                            except ImportError as import_err:
                                                print(f"[UPLOAD] ❌ OCR libraries not installed!")
                                                print(f"[UPLOAD] Missing: {import_err}")
                                                print(f"[UPLOAD] Install with: pip install pdf2image pytesseract Pillow")
                                                print(f"[UPLOAD] Also install Tesseract OCR engine:")
                                                print(f"[UPLOAD]   macOS: brew install tesseract")
                                                print(f"[UPLOAD]   Linux: sudo apt-get install tesseract-ocr")
                                                print(f"[UPLOAD]   Windows: https://github.com/UB-Mannheim/tesseract/wiki")
                                                # Don't fail completely - will add error message below
                                            except Exception as ocr_error:
                                                print(f"[UPLOAD] ❌ OCR failed with error: {ocr_error}")
                                                print(f"[UPLOAD] Error type: {type(ocr_error).__name__}")
                                                import traceback
                                                print(f"[UPLOAD] Traceback: {traceback.format_exc()}")
                                            
                                            # If OCR failed and we have no text, log error but don't add error message as content
                                            if not ocr_success and len(extracted_text.strip()) < 100:
                                                print(f"[UPLOAD] ❌ CRITICAL: OCR failed and no text extracted. PDF cannot be processed.")
                                                print(f"[UPLOAD] This is a scanned PDF that requires OCR, but OCR is not working.")
                                                # Don't add error message as content - it confuses the transformer
                                                # Instead, skip this file or add a minimal placeholder
                                                extracted_text = f"[SCANNED PDF - OCR FAILED]\n\nUnable to extract text from this scanned PDF. OCR processing failed.\n\nFilename: {file_name}\nPages: {len(pdf_reader.pages)}\n\nPlease check OCR installation and try again."
                                        
                                        # Add to file_contents - either the extracted text or an error
                                        if extracted_text and len(extracted_text.strip()) > 50 and not extracted_text.startswith("[SCANNED PDF"):
                                            # Success: we have meaningful extracted text
                                            final_text = extracted_text[:100000]
                                            if len(extracted_text) > 100000:
                                                final_text += f"\n\n[Document truncated - {len(extracted_text)} total chars]"
                                            
                                            file_contents.append(f"[PDF File: {file_name}]\n{final_text}")
                                            print(f"[UPLOAD] ✅ Final extracted {len(final_text)} chars from PDF: {file_name}")
                                        else:
                                            # OCR failed - add error so transformer knows a file was uploaded but failed
                                            print(f"[UPLOAD] ❌ OCR failed for PDF {file_name}")
                                            file_contents.append(f"[PDF File: {file_name}]\n[ERROR: This is a scanned/image-based PDF. OCR text extraction failed. Please install OCR dependencies: pip install pdf2image pytesseract && brew install tesseract poppler]")
                                        
                                    except Exception as e:
                                        print(f"[UPLOAD] Failed to parse PDF {file_name}: {e}")
                                        file_contents.append(f"[PDF File: {file_name}]\n[Error parsing PDF: {str(e)}]")
                                
                                # Parse DOCX files
                                elif file_content.startswith("__DOCX_BASE64__"):
                                    try:
                                        import base64
                                        from io import BytesIO
                                        from docx import Document
                                        
                                        docx_base64 = file_content[15:]  # Remove prefix
                                        docx_bytes = base64.b64decode(docx_base64)
                                        doc = Document(BytesIO(docx_bytes))
                                        
                                        # Extract text from paragraphs and tables
                                        text_parts = []
                                        for para in doc.paragraphs:
                                            if para.text.strip():
                                                text_parts.append(para.text)
                                        
                                        for table in doc.tables:
                                            for row in table.rows:
                                                row_text = " | ".join(cell.text for cell in row.cells)
                                                text_parts.append(row_text)
                                        
                                        extracted_text = "\n".join(text_parts)
                                        file_contents.append(f"[Word File: {file_name}]\n{extracted_text[:100000]}")
                                        print(f"[UPLOAD] Extracted {len(extracted_text)} chars from DOCX: {file_name}")
                                    except Exception as e:
                                        print(f"[UPLOAD] Failed to parse DOCX {file_name}: {e}")
                                        file_contents.append(f"[Word File: {file_name}]\n[Error parsing DOCX: {str(e)}]")
                                
                                # Plain text content
                                else:
                                    file_contents.append(f"[File: {file_name}]\n{file_content[:100000]}")
                        
                        if file_contents:
                            context["uploaded_file_content"] = "\n\n".join(file_contents)
                            context["user_message"] = f"{user_message}\n\nUploaded files:\n{context['uploaded_file_content']}"
                            print(f"[UPLOAD] Set uploaded_file_content with {len(context['uploaded_file_content'])} chars")
                            print(f"[UPLOAD] Content preview: {context['uploaded_file_content'][:500]}...")
                        else:
                            # No content extracted from any files - set error message
                            print(f"[UPLOAD] ⚠️ WARNING: {len(uploaded_files)} files uploaded but no content extracted!")
                            file_names = [f.get("name", "unknown") for f in uploaded_files]
                            context["uploaded_file_content"] = f"[UPLOAD ERROR: Files uploaded ({', '.join(file_names)}) but content extraction failed. If these are scanned PDFs, OCR may not be installed or working. Install: pip install pdf2image pytesseract && brew install tesseract poppler]"
                            context["user_message"] = f"{user_message}\n\n{context['uploaded_file_content']}"
                    else:
                        print(f"[UPLOAD] No files in uploadedFiles array")
                    
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "input",
                            "content": f"Uploaded and parsed {len(uploaded_files)} file(s)",
                        }
                    })
                else:
                    # Prompt node - use promptText if available, otherwise user_message
                    prompt_text = node_data.get("promptText", user_message)
                    if prompt_text:
                        context["user_message"] = prompt_text
                    
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "input",
                            "content": prompt_text or user_message,
                        }
                    })
                continue
            
            if node_type in OUTPUT_NODE_TYPES:
                executed_nodes.add(node_id)
                final_content = context.get("final_answer", "")
                
                # For spreadsheet output, format as CSV/table data
                if node_type == "spreadsheet":
                    # Store spreadsheet flag for final output
                    context["output_format"] = "spreadsheet"
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "spreadsheet_output",
                            "content": final_content,
                            "format": "spreadsheet",
                        }
                    })
                else:
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "output",
                            "content": final_content,
                        }
                    })
                continue
            
            # Check dependencies (already computed above)
            missing_deps = [
                dep for dep in dependencies
                if dep not in executed_nodes and dep not in excluded_nodes
            ]
            
            if missing_deps:
                debugger.log_node_skipped(node_id, f"Missing dependencies: {missing_deps}")
                continue
            
            # === BRANCH ROUTING LOGIC ===
            # A node should only execute if at least one of its upstream dependencies was executed
            # If ALL dependencies were excluded, this node should also be excluded
            should_execute = True
            
            workflow_logger.debug(f"Branch routing check for {node_id}:")
            workflow_logger.debug(f"  Dependencies: {dependencies}")
            workflow_logger.debug(f"  Executed nodes: {executed_nodes}")
            workflow_logger.debug(f"  Excluded nodes: {excluded_nodes}")
            
            if dependencies:
                # Check if any dependency was actually executed (not excluded)
                executed_deps = [dep for dep in dependencies if dep in executed_nodes and dep not in excluded_nodes]
                excluded_deps = [dep for dep in dependencies if dep in excluded_nodes]
                
                workflow_logger.debug(f"  Executed dependencies: {executed_deps}")
                workflow_logger.debug(f"  Excluded dependencies: {excluded_deps}")
                
                has_executed_dependency = len(executed_deps) > 0
                
                if not has_executed_dependency:
                    # All our dependencies were excluded - we should be excluded too
                    should_execute = False
                    excluded_nodes.add(node_id)
                    
                    debugger.log_branch_decision(
                        node_id, node_type, "EXCLUDE",
                        f"All dependencies excluded: {excluded_deps}",
                        {"executed_deps": executed_deps, "excluded_deps": excluded_deps}
                    )
                    debugger.log_node_excluded(node_id, node_type, "All upstream dependencies were excluded")
                    
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "exclude",
                            "content": "Excluded (upstream path not taken)",
                            "excluded": True,
                        }
                    })
                    continue
                else:
                    debugger.log_branch_decision(
                        node_id, node_type, "EXECUTE",
                        f"Has executed dependencies: {executed_deps}",
                        {"executed_deps": executed_deps, "excluded_deps": excluded_deps}
                    )
            
            # === ORCHESTRATOR BRANCH ROUTING ===
            # When orchestrator selects specific tools, ONLY those paths should execute.
            # All other paths from orchestrator should be excluded.
            
            # Check if this node is directly connected to an orchestrator
            orchestrator_parent = None
            for edge in valid_edges:
                if edge["target"] == node_id and edge["source"].startswith("orchestrator"):
                    orchestrator_parent = edge["source"]
                    break
            
            if orchestrator_parent:
                tools_to_execute = context.get("orchestrator_result", {}).get("tools_to_execute", [])
                
                workflow_logger.info(f"ORCHESTRATOR CHILD check for {node_id} ({node_type}):")
                workflow_logger.info(f"  Orchestrator parent: {orchestrator_parent}")
                workflow_logger.info(f"  Tools selected by orchestrator: {tools_to_execute}")
                
                # Determine if this node should execute based on orchestrator decision
                should_execute_based_on_orchestrator = False
                
                if "image_generator" in tools_to_execute:
                    # Image path was selected - only image_generator should execute
                    if node_type == "image_generator":
                        should_execute_based_on_orchestrator = True
                        workflow_logger.info(f"  -> EXECUTE: This is the selected image_generator path")
                    else:
                        # This is an alternative path (e.g., semantic_search) - exclude it
                        should_execute_based_on_orchestrator = False
                        workflow_logger.info(f"  -> EXCLUDE: image_generator selected, this path ({node_type}) excluded")
                else:
                    # No image_generator selected - default text path executes
                    if node_type == "image_generator":
                        should_execute_based_on_orchestrator = False
                        workflow_logger.info(f"  -> EXCLUDE: image_generator not selected")
                    else:
                        # Default path (semantic_search, etc.) should execute
                        should_execute_based_on_orchestrator = True
                        workflow_logger.info(f"  -> EXECUTE: Default path (no image_generator selected)")
                
                if not should_execute_based_on_orchestrator:
                    should_execute = False
                    excluded_nodes.add(node_id)
                    
                    debugger.log_branch_decision(
                        node_id, node_type, "EXCLUDE",
                        f"Orchestrator routing: Selected tools={tools_to_execute}, this node type={node_type}",
                        {"tools_to_execute": tools_to_execute, "node_type": node_type}
                    )
                    debugger.log_node_excluded(node_id, node_type, f"Not on selected orchestrator path")
                    
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "exclude",
                            "content": f"Excluded (orchestrator selected: {tools_to_execute or 'default text path'})",
                            "excluded": True,
                        }
                    })
                    continue
                else:
                    debugger.log_branch_decision(
                        node_id, node_type, "EXECUTE",
                        f"On selected orchestrator path: tools={tools_to_execute}",
                        {"tools_to_execute": tools_to_execute}
                    )
            
            # === SUPERVISOR PATH ROUTING ===
            # When supervisor explicitly selects a path (IMAGE_GENERATOR or SEMANTIC_SEARCH),
            # ONLY that path should execute. ALL OTHER NODES ON DIFFERENT PATHS should be excluded.
            # This applies to ALL nodes, not just direct children of supervisor.
            
            # Parse supervisor guidance to find WORKFLOW PATH (check for any supervisor in workflow)
            supervisor_guidance = context.get("supervisor_guidance", "")
            supervisor_plan = context.get("supervisor_plan", "")
            
            # Check both guidance and plan for workflow path
            workflow_path_text = f"{supervisor_guidance}\n{supervisor_plan}"
            
            # Extract the workflow path selected by supervisor
            selected_path = None
            if "WORKFLOW PATH:" in workflow_path_text.upper():
                for line in workflow_path_text.split("\n"):
                    if "WORKFLOW PATH:" in line.upper():
                        path_line = line.upper()
                        if "IMAGE_GENERATOR" in path_line or "IMAGE GENERATOR" in path_line:
                            selected_path = "IMAGE_GENERATOR"
                            break
                        elif "SEMANTIC_SEARCH" in path_line or "SEMANTIC SEARCH" in path_line:
                            selected_path = "SEMANTIC_SEARCH"
                            break
            
            # Apply supervisor path routing to ALL relevant nodes (not just supervisor children)
            if selected_path:
                workflow_logger.info(f"SUPERVISOR PATH check for {node_id} ({node_type}):")
                workflow_logger.info(f"  Supervisor selected path: {selected_path}")
                
                # Define which node types belong to which path
                image_path_nodes = {"image_generator"}
                text_path_nodes = {"semantic_search", "sampler", "synthesis"}
                neutral_nodes = {"prompt", "supervisor", "orchestrator", "response", "spreadsheet", "transformer"}
                
                # Determine if this node should execute based on supervisor path selection
                should_execute_based_on_supervisor = True
                exclude_reason = None
                
                if selected_path == "IMAGE_GENERATOR":
                    # Image path was selected
                    if node_type in text_path_nodes:
                        # This is a text/search path node - EXCLUDE IT
                        should_execute_based_on_supervisor = False
                        exclude_reason = f"IMAGE_GENERATOR path selected, {node_type} is on text path"
                        workflow_logger.info(f"  -> EXCLUDE: {exclude_reason}")
                    elif node_type in image_path_nodes:
                        workflow_logger.info(f"  -> EXECUTE: This is the selected IMAGE_GENERATOR path")
                    elif node_type in neutral_nodes:
                        workflow_logger.info(f"  -> EXECUTE: Neutral node type ({node_type})")
                    else:
                        workflow_logger.info(f"  -> EXECUTE: Unknown node type ({node_type}), allowing")
                        
                elif selected_path == "SEMANTIC_SEARCH":
                    # Text/search path was selected
                    if node_type in image_path_nodes:
                        # This is an image path node - EXCLUDE IT
                        should_execute_based_on_supervisor = False
                        exclude_reason = f"SEMANTIC_SEARCH path selected, {node_type} is on image path"
                        workflow_logger.info(f"  -> EXCLUDE: {exclude_reason}")
                    elif node_type in text_path_nodes:
                        workflow_logger.info(f"  -> EXECUTE: This is the selected SEMANTIC_SEARCH path")
                    elif node_type in neutral_nodes:
                        workflow_logger.info(f"  -> EXECUTE: Neutral node type ({node_type})")
                    else:
                        workflow_logger.info(f"  -> EXECUTE: Unknown node type ({node_type}), allowing")
                
                if not should_execute_based_on_supervisor:
                    should_execute = False
                    excluded_nodes.add(node_id)
                    
                    debugger.log_branch_decision(
                        node_id, node_type, "EXCLUDE",
                        f"Supervisor routing: {exclude_reason}",
                        {"selected_path": selected_path, "node_type": node_type}
                    )
                    debugger.log_node_excluded(node_id, node_type, f"Not on selected supervisor path: {selected_path}")
                    
                    yield _sse_event("agent_complete", {
                        "agent": node_id,
                        "step": {
                            "agent": node_type,
                            "model": "none",
                            "action": "exclude",
                            "content": f"Excluded (supervisor selected: {selected_path})",
                            "excluded": True,
                        }
                    })
                    continue
                else:
                    debugger.log_branch_decision(
                        node_id, node_type, "EXECUTE",
                        f"On selected supervisor path: {selected_path}",
                        {"selected_path": selected_path}
                    )
            
            if not should_execute:
                continue
            
            # Execute the agent
            yield _sse_event("agent_start", {"agent": node_id, "status": "working"})
            
            # Use context user_message (which may include uploaded file content)
            effective_message = context.get("user_message", user_message)
            
            result = await _execute_agent(
                node_type=node_type,
                user_message=effective_message,
                context=context,
                settings=node_settings,
                llm_client=llm_client,
                small_model=small_model,
                large_model=large_model,
                valid_edges=valid_edges,
                reachable_nodes=reachable_nodes,
                node_map=node_map,
            )
            
            if result:
                # Update context with agent's results
                workflow_logger.debug(f"Context updates from {node_id}:")
                for key, value in result.context_updates.items():
                    debugger.log_context_update(key, value, node_id)
                    
                    if key == "context_snippets":
                        context["context_snippets"].extend(value)
                    elif key == "images":
                        context["tool_outputs"]["images"].extend(value)
                    elif key == "docs":
                        context["docs"].extend(value)
                    else:
                        context[key] = value
                    
                    # Special logging for orchestrator decisions
                    if key == "orchestrator_result":
                        tools = value.get("tools_to_execute", [])
                        reasoning = value.get("reasoning", "")
                        debugger.log_orchestrator_decision(
                            tools,
                            context.get("available_tools", []),
                            reasoning
                        )
                
                # Record step
                step = {
                    "agent": result.agent,
                    "model": result.model,
                    "action": result.action,
                    "content": result.content,
                    **result.metadata,
                }
                steps.append(step)
                executed_nodes.add(node_id)
                
                debugger.log_node_execution(node_id, node_type, result.action, result.content)
                
                yield _sse_event("agent_complete", {"agent": node_id, "step": step})
            else:
                executed_nodes.add(node_id)
                yield _sse_event("agent_complete", {
                    "agent": node_id,
                    "step": {
                        "agent": node_type,
                        "model": "none",
                        "action": "skip",
                        "content": "Skipped",
                    }
                })
        
        # Determine final answer
        final_answer = context.get("final_answer", "")
        if not final_answer and context.get("context_snippets"):
            final_answer = "\n\n".join(context["context_snippets"])
        if not final_answer:
            final_answer = "No output generated."
        
        # Calculate metrics
        workflow_latency = round((time.time() - workflow_start) * 1000, 2)
        
        # Log execution summary
        debugger.log_execution_summary(
            executed_nodes,
            excluded_nodes,
            len(reachable_nodes),
            workflow_latency
        )
        
        # Prepare final tool outputs
        final_tool_outputs = {
            "images": [
                {
                    "prompt": img.get("prompt"),
                    "style": img.get("style"),
                    "url": img.get("url"),
                    "has_data": bool(img.get("url")),
                }
                for img in context["tool_outputs"].get("images", [])
            ],
            "calculations": context["tool_outputs"].get("calculations", []),
            "web_results": context["tool_outputs"].get("web_results", []),
            "docs": context.get("docs", []),
        }
        
        # Check if spreadsheet output was requested
        output_format = context.get("output_format", "text")
        
        workflow_logger.info(f"Final output format: {output_format}")
        workflow_logger.info(f"Final answer length: {len(final_answer)} chars")
        
        yield _sse_event("done", {
            "answer": final_answer,
            "tool_outputs": final_tool_outputs,
            "trace": {"steps": steps},
            "latency_ms": workflow_latency,
            "output_format": output_format,
        })
        
    except Exception as exc:
        debugger.log_error(f"Workflow execution failed", exc)
        yield _sse_event("error", {"message": str(exc)})


async def _execute_agent(
    node_type: str,
    user_message: str,
    context: Dict[str, Any],
    settings: Dict[str, Any],
    llm_client: LLMClientProtocol,
    small_model: str,
    large_model: str,
    valid_edges: List[Dict[str, str]],
    reachable_nodes: Set[str],
    node_map: Dict[str, Any],
) -> Optional[AgentResult]:
    """
    Execute a single agent based on node type.
    
    Args:
        node_type: Type of node to execute
        user_message: Original user query
        context: Shared execution context
        settings: Node-specific settings
        llm_client: LLM client for completions
        small_model: Small model name
        large_model: Large model name
        valid_edges: Valid workflow edges
        reachable_nodes: Set of reachable node IDs
        
    Returns:
        AgentResult or None if agent not found
    """
    # Get agent class from registry
    agent_class = AGENT_REGISTRY.get(node_type)
    
    if not agent_class:
        print(f"[WORKFLOW] Unknown agent type: {node_type}")
        return None
    
    # Create agent instance
    if node_type == "semantic_search":
        agent = agent_class(llm_client, retrieval)
    else:
        agent = agent_class(llm_client)
    
    # Determine which model to use
    if agent.default_model == "large":
        model = large_model
    elif agent.default_model == "embedding":
        model = None  # Semantic search uses embeddings
    else:
        model = small_model
    
    # Add downstream node TYPES to context for supervisor (so it understands workflow structure)
    if node_type == "supervisor":
        # Get the node types, not IDs - supervisor needs to know what's in the workflow
        downstream_types = set()
        for nid in reachable_nodes:
            # Extract type from node ID (format: "type-timestamp")
            if "-" in nid:
                ntype = nid.rsplit("-", 1)[0]
                downstream_types.add(ntype)
            else:
                downstream_types.add(nid)
        context["downstream_nodes"] = list(downstream_types)
        print(f"[SUPERVISOR] Downstream node types: {downstream_types}")
    
    # Add available tools to context for orchestrator
    if node_type == "orchestrator":
        available_tools = []
        for node_id in reachable_nodes:
            # Look up the node in node_map to get its actual type
            node = node_map.get(node_id)
            if node:
                node_data = node.get("data", {})
                other_node_type = node_data.get("nodeType", "")
                
                # Check if this is a tool node that should be available
                if other_node_type == "image_generator" and "image_generator" not in available_tools:
                    available_tools.append("image_generator")
                elif other_node_type == "web_search" and "web_search" not in available_tools:
                    available_tools.append("web_search")
        
        workflow_logger.debug(f"Orchestrator available tools detection:")
        workflow_logger.debug(f"  Reachable nodes: {reachable_nodes}")
        for node_id in reachable_nodes:
            node = node_map.get(node_id)
            if node:
                node_data = node.get("data", {})
                other_node_type = node_data.get("nodeType", "")
                workflow_logger.debug(f"    {node_id} -> {other_node_type}")
        workflow_logger.debug(f"  Detected available tools: {available_tools}")
        
        context["available_tools"] = available_tools
    
    # Execute agent
    result = await agent.execute(
        user_message=user_message,
        context=context,
        settings=settings,
        model=model,
    )
    
    return result

